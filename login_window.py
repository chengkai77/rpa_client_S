# -*- coding: utf-8 -*-
import time
import tkinter as tk
import base64
import os
import win32gui
import win32print
import win32con
from PIL import Image, ImageTk, ImageSequence
from docx import Document
from tkinter import *
# from mttkinter import mtTkinter as tk
import tkinter.messagebox


def create_pic():
    base64_data = b'AAABAAEAJDAAAAEAIACoHAAAFgAAACgAAAAkAAAAYAAAAAEAIAAAAAAAABsAABAnAAAQJwAAAAAAAAAAAAAAAAAAAAAAAAAAAAB+fn4AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB8fHwAfn5+AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQB9fX0AfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19AH19fQF9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0DfX19AH19fQB9fX0AfX19AH19fQN9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0IfX19CH19fQh9fX0BfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19EH19fYl9fX29fX19vX19fb19fX29fX19vX19fb19fX29fX19vX19fb59fX2ffX19I319fQB9fX0AfX19I319fZ99fX2+fX19vX19fb19fX29fX19vX19fb19fX29fX19vX19fb19fX2JfX19EH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19SH19ffx9fX3/fX19/319ff99fX3/fX19/319ff99fX3/fX19/319ff99fX3/fX19eH19fQB9fX0AfX19eH19ff99fX3/fX19/319ff99fX3/fX19/319ff99fX3/fX19/319ff99fX38fX19SH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19J319fc19fX30fX198319ffR/f3/9gICA/4CAgP+AgID/gICA/4CAgP+AgID/gICAhYCAgACAgIAAgICAhYCAgP+AgID/gICA/4CAgP+AgID/gICA/39/f/19fX30fX198319ffR9fX3NfX19J319fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19AH19fR19fX06fHx8OYiIiEeqqqrZra2t/62trf6tra3+ra2t/q2trf6tra3/rKysh66urgCurq4ArKysh62trf+tra3+ra2t/q2trf6tra3+ra2t/6qqqtmIiIhHfHx8OX19fTp9fX0dfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9fX0AfX19AH19fQB9fX0Ai4uLAMfHxw67u7vOu7u7/7u7u/+7u7v/u7u7/7u7u/+7u7v/urq6iL6+vgC+vr4Aurq6iLu7u/+7u7v/u7u7/7u7u/+7u7v/u7u7/7u7u87Hx8cOi4uLAH19fQB9fX0AfX19AH19fQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAuLi4ALi4uBG6urrOurq6/7q6uv+6urr/urq6/7q6uv+6urr/ubm5iL29vQC9vb0Aubm5iLq6uv+6urr/urq6/7q6uv+6urr/urq6/7q6us64uLgRuLi4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAnp6eAJycnACcnJwAnJycAJycnACdnZ0AuLi4ALi4uBG6urrOurq6/7q6uv+6urr/urq6/7q6uv+6urr/ubm5iL29vQC9vb0Aubm5iLq6uv+6urr/urq6/7q6uv+6urr/urq6/7q6us64uLgRuLi4AJ2dnQCcnJwAnJycAJycnACdnZ0Am5ubAAAAAACcnJwAnJycAJycnACcnJwAnJycAJycnACcnJwAt7e3ALi4uBG6urrOurq6/7q6uv+6urr/urq6/7q6uv+6urr/ubm5iL29vQC9vb0Aubm5iLq6uv+6urr/urq6/7q6uv+6urr/urq6/7q6us64uLgRt7e3AJycnACcnJwAnJycAJycnACcnJwAnJycAJycnACbm5sAnJycAJycnC2cnJyOnJycjZycnCycnJwApqamAKamphCnp6fNp6en/6enp/+np6f/p6en/6enp/+np6f/p6enhqenpwCnp6cAp6enhqenp/+np6f/p6en/6enp/+np6f/p6en/6enp82mpqYQp6enAJycnACcnJwlnJyciJycnJOcnJw1nJycAJycnACcnJwAnJycGZycnMecnJz/nJyc/5ycnMScnJwVnJycAJubmxCbm5vMm5ub/5ubm/+bm5v/m5ub/5ubm/+bm5v/nJychJubmwCbm5sAnJychJubm/+bm5v/m5ub/5ubm/+bm5v/m5ub/5ubm8ybm5sQnJycAJycnBCcnJy5nJyc/5ycnP+cnJzRnJycH5ycnACcnJwAnJycRZycnPmcnJz/nJyc/5ycnPOcnJw3nJycAJycnBCcnJzMnJyc/5ycnP+cnJz/nJyc/5ycnP+cnJz/nJychJycnACcnJwAnJychJycnP+cnJz/nJyc/5ycnP+cnJz/nJyc/5ycnMycnJwQnJycAJycnDWcnJzxnJyc/5ycnP+cnJz6nJycSJycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnPScnJw7nJycAJycnBCcnJzMnJyc/5ycnP+cnJz/nJyc/5ycnP+cnJz/nJychJycnACcnJwAnJychJycnP+cnJz/nJyc/5ycnP+cnJz/nJyc/5ycnMycnJwQnJycAJycnDucnJz0nJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnPOcnJw7nJycAJ2dnQ+cnJzMnJyc/5ycnP+cnJz/nJyc/5ycnP+cnJz/nJychJORjgCTkY4AnJychJycnP+cnJz/nJyc/5ycnP+cnJz/nJyc/5ycnMydnZ0PnJycAJycnDucnJzznJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnPOcnJ06eHFlAJSTkBGZmZjNmpmZ/5qZmf6amZn+mpmZ/pqZmf6amZn/mJeWiCoOAAUqDgAFmJeWiJqZmf+amZn+mpmZ/pqZmf6amZn+mpmZ/5mZmM2Uk5AReHFlAJycnTqcnJzznJyc/5ycnP+cnJz6nJycS5ycnACbm5sAnJycS5SUlPqIiIj/iIiI/5WVlfOamZg9Qi4PSUYzFqpYSDHwW003/VtNN/1bTTf9W003/VtNN/1cTTf+UkEp3UQwE7pEMBO6UkEp3VxNN/5bTTf9W003/VtNN/1bTTf9W003/VhIMfBGMxaqQi4PSZqZmD2VlZXziIiI/4iIiP+UlJT6nJycS5ubmwCMjIwAjY2NS4GBgfp8fHz/fHx8/4ODg/FlWUmNQi4P7EIuEP9BLQ7/QSwO/0EsDv9BLA7/QSwO/0EsDv9BLA7/QS0P/0IuEP9CLhD/QS0P/0EsDv9BLA7/QSwO/0EsDv9BLA7/QSwO/0EtDv9CLhD/Qi4P7GVZSYyDg4PxfHx8/3x8fP+BgYH6jY2NS4yMjAB+fn4AgICAS319ffp9fX3/fX19/3x8e/lTRC7sQS0O/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/QS0P/0EtDv9BLQ7/QS0P/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/QS0O/1NELux8fHv5fX19/319ff99fX36gICAS35+fgCOjo4Aj4+PS4KCgvp8fHz/fX19/4B/fv9URS7/QS0O/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9BLQ//Szkd/1VEK/9VRCr/Sjcb/0EtD/9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/QS0O/1RFLv+Af37/fX19/3x8fP+CgoL6j4+PS46OjgCcnJwAnJycS5WVlfqKior/i4uL/5KSkP9XRzD+QCwN/0IuD/9CLg//Qi4P/0IuD/9CLg//Qi0P/0AsDv9cTDT/n5qT/7CurP+wrqv/m5aO/1dHLv9ALA7/Qi4P/0IuD/9CLg//Qi4P/0IuD/9CLg//QCwN/1dHMP6SkpD/i4uL/4qKiv+VlZX6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyd/5mYlv9cTTf+RzQX/0g1Gf9INRn/SDUZ/0g1Gf9INRn/SDUZ/2BROv+lop3/vLy9/7u7u/+7u7v/vLy9/6Gdl/9cTDX/SDUZ/0g1Gf9INRn/SDUZ/0g1Gf9INRn/RzQX/1xNN/6ZmJb/nJyd/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/5ycnP6fnJf9oJyW/6Cclv+gnJb/oJyW/6Cclv+gnJb/oZ2X/7CurP+7u7v/urq6/7q6uv+6urr/urq6/7u7u/+vrar/oZyW/6Cclv+gnJb/oJyW/6Cclv+gnJb/oJyW/5+cl/2cnJz+nJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqcnJz/nJyc/52dnf61tbX8vLy9/7y8vP+8vLz/vLy8/7y8vP+8vLz/vLy8/7u7u/+6urr/urq6/7q6uv+6urr/urq6/7q6uv+7u7v/u7y8/7u8vP+7vLz/u7y8/7u8vP+8vLz/vLy9/7W1tfydnZ3+nJyc/5ycnP+cnJz6nJycS5ycnACcnJwAnJycS5ycnPqbm5v/m5ub/56env6zs7P8urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+9vb3/vb29/729vf+9vb3/vb29/729vf+6urr/urq6/7Ozs/yenp7+m5ub/5ubm/+cnJz6nJycS5ycnACbm5sAm5ubS5GRkfqEhIT/hYWF/5SUlP6zs7P8urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/8DAwP/d3d3/4uLi/+Hh4f/h4eH/4uLi/9nZ2f+8vLz/urq6/7Ozs/yUlJT+hYWF/4SEhP+RkZH6m5ubS5ubmwCJiYkAioqKS4CAgPp8fHz/fHx8/4ODg/6wsLD8u7u7/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/76+vv/T09P/1tbW/9XV1f/V1dX/1tbW/9DQ0P+8vLz/u7u7/7CwsPyDg4P+fHx8/3x8fP+AgID6ioqKS4mJiQB+fn4Afn5+Rn19ffl9fX3/fX19/4CAgP6tra38u7u7/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/u7u7/62trfyAgID+fX19/319ff99fX35fn5+Rn5+fgB/f38AgICAGn5+fsd9fX3/fX19/4iIiP+ysrL/u7u7/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/urq6/7q6uv+6urr/u7u7/7Kysv+IiIj/fX19/319ff9+fn7GgICAGn9/fwBra2sAe3t7AICAgCd+fn55g4ODiJeXl4a0tLSFurq6h7q6uoa6urqGurq6hrq6uoa6urqGurq6iK+vr9SsrKz8rKys+qysrPqsrKz6rKys+qysrPyvr6/Uurq6iLq6uoa6urqGurq6hrq6uoa6urqGurq6h7S0tIWXl5eGgoKCiH5+fnmAgIAne3t7AGxsbACGhoYAbGxsAH9/fwB7e3sAc3NzAIuLiwC0tLQAurq6ALq6ugC6uroAurq6ALq6ugC4uLgA////AJubm6qcnJz/nJyc/5ycnP+cnJz/nJyc/5ycnP+bm5uq////ALi4uAC6uroAurq6ALq6ugC6uroAurq6ALS0tACLi4sAc3NzAHt7ewB/f38AbW1tAIaGhgAAAAAAqqqqAIiIiACJiYkAlJSUAJ+fnwC7u7sAurq6ALq6ugC8vL0AiIB0ADsmBgBBLQ8AAAAAAJ6en6qdnZ7/nZ6e/52env+dnp7/nZ6e/52dnv+enp+qAAAAAEEtDwA7JgYAiIB0ALy8vQC6uroAurq6ALu7uwCfn58AlJSUAIiIiACIiIgAqqqqAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroAurq6Arq6ujm8vb14ioJ2hEQxE4VFMRSFRTIVh2ZaSNZvZlf/b2VW/m9lVv5vZVb+b2VW/m9mV/9mWkjWRTIVh0UxFIVEMROFioJ2hLy9vXi6uro6urq6Arq6ugC6uroAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroAurq6Xbq6uuu8vL3/iYF1/0IuD/9CLhD/Qi4Q/0EtD/9BLA7/QS0O/0EtDv9BLQ7/QS0O/0EsDv9BLQ//Qi4Q/0IuEP9CLg//iYF1/7y8vf+6urrrurq6Xbq6ugC6uroAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAJKSkgCcnJwAnJycALGxsQC7u7sVurq60Lq6uv+8vL3/iYF1/0ItD/9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLhD/Qi4Q/0IuEP9CLQ//iYF1/7y8vf+6urr/urq60Lu7uxWxsbEAnJycAJycnACSkpIAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwAmZmZAJCQkAe1tbU0urq667q6uv+8vL3/iYF1/0ItD/9CLhD/QS0P/0QxE/9FMRT/QS0P/0IuEP9CLhD/QS0P/0UxFP9EMRP/QS0P/0IuEP9CLQ//iYF1/7y8vf+6urr/urq667W1tTSQkJAHmpqaAJycnACcnJwAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwEnJycbZycnMCioqLNuLi4+bq6uv+8vL3/iYF1/0ItD/9BLQ//VUQp/6ujmP+wqZ//W0ox/0EtDv9BLQ7/W0ox/7Cpn/+ro5j/VUQp/0EtD/9CLQ//iYF1/7y8vf+6urr/uLi4+aKios2cnJzAnJycbZycnAScnJwAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwfnJyc4JycnP+hoaH/uLi4/rq6uv+8vL3/iYF1/0ItD/8/Kwz/hHdl/+zt7f/u7/D/kIZ2/z8rDf9AKw3/kIZ2/+7v8P/s7e3/g3dl/z8rDP9CLQ//iYF1/7y8vf+6urr/uLi4/qGhof+cnJz/nJyc4JycnB+cnJwAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwmnJyc55ycnP+hoaH9uLi4/rq6uv+8vL3/iYF1/0EtD/9ALA7/ZlY//87Lxv/T0Mz/b2BK/0AsDv9ALA7/b2BK/9PQzP/Oy8b/ZlY//0AsDv9BLQ//iYF1/7y8vf+6urr/uLi4/qGhof2cnJz/nJyc6Jubmyebm5sAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwZnJyc1pycnP+hoaH/uLi4/rq6uv+8vL3/jod8/0IuEP9CLhD/Qi4Q/1ZFKv9YRy3/Qy8R/0IuEP9CLhD/Qy8R/1hHLf9WRCr/Qi4Q/0IuEP9CLhD/jod8/7y8vf+6urr/uLi4/qGhof+cnJz/mpqa55KSkiaSkpIAAAAAAAAAAAAAAAAAAAAAAJycnACcnJwBnJycSZubm5OkpKSpubm59rq6uv+7u7z/qKWg/1JBJ/9AKw3/QS0P/0AsDf9ALA3/QS0P/0EtD/9BLQ//QS0P/0AsDf9ALA3/QS0P/0ArDf9SQSf/qKWg/7u7vP+6urr/ubm59qSkpKeUlJTGh4eH535+fiV/f38AAAAAAAAAAAAAAAAAAAAAAJycnACcnJwAnJycAK+vrwC7u7srurq667q6uv+6urr/urq6/5eRiP9cTTb/TDof/0w5Hv9MOR7/TDke/0w5Hv9MOR7/TDke/0w5Hv9MOR7/TDof/1xNNf+XkYj/urq6/7q6uv+6urr/urq668HBwSd7e3tvfX196H19fSV9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAACcnJwAnJycALq6ugC6urokurq65bq6uv+6urr/urq6/7u7u/+ysa//p6Sf/6ainf+mop3/pqKd/6ainf+mop3/pqKd/6ainf+mop3/p6Sf/7Kxr/+7u7v/urq6/7q6uv+6urr/urq65cLCwiB8fHxwfX196H19fSV9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroHurq6pbq6uv+6urr/urq6/7q6uv+7u7v/u7u8/7u7vP+7u7z/u7u8/7u7vP+7u7z/u7u8/7u7vP+7u7z/u7u8/7u7u/+6urr/urq6/7q6uv+6urr/urq6pf///wN9fX1yfX196H19fSV9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALq6ugC6uroAurq6Ibq6uqG6urrfurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6uuW6urrlurq65bq6ut+6urqhurq6IZCQkAB9fX10fX197H19fSZ9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALu7uwC6uroAurq6ALq6ugW6uroeurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uiS6urokurq6JLq6uh66uroFurq6AHx8fAB9fX1UfX19qn19fRt9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6ALq6ugC6uroAurq6AH19fQB9fX0DfX19B319fQF9fX0AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH19fQB9fX0AfX19AH19fQB9fX0AAAAAAAAAAADgAAAAcAAAAOAAAABwAAAA4AAAAHAAAADgAAAAcAAAAOAAAABwAAAA4AAAAHAAAADgAAAAcAAAAP4AAAfwAAAAgAAAABAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAIAAAAAQAAAA+AAAAfAAAAD4AAAB8AAAAMAAAAAwAAAAwAAAADAAAADAAAAAMAAAAMAAAAAwAAAAwAAAADAAAADAAAAAMAAAAMAAAAAwAAAAwAAAADAAAADgAAAAMAAAAPgAAAAwAAAA+AAAADAAAAD4AAAAMAAAAPwAAAAwAAAA////+DAAAAA='
    img_data = base64.b64decode(base64_data)
    # 注意：如果是"data:image/jpg:base64,"，那你保存的就要以png格式，如果是"data:image/png:base64,"那你保存的时候就以jpg格式。
    with open('window_title.png', 'wb') as f:
        f.write(img_data)


def delete_pic():
    try:
        os.remove("window_title.png")
    except Exception as e:
        print(e)


def info_remainder(info):
    create_pic()
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    root.iconbitmap("window_title.png")
    tk.messagebox.showinfo("ChocLead", info)
    delete_pic()


def info_window(title, info):
    create_pic()
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    root.iconbitmap("window_title.png")
    tk.messagebox.showinfo(title, info)
    delete_pic()


def error_remainder(info):
    create_pic()
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    root.iconbitmap("window_title.png")
    tk.messagebox.showerror("ChocLead", info)
    delete_pic()


# robot connect loading image
def pick():
    global a, flag
    root = tk.Tk()
    # root.attributes('-toolwindow', True,
    #                 '-alpha', 0.2)
    root.overrideredirect(True)
    width = 300
    height = 300
    screenwidth = root.winfo_screenwidth()
    screenheight = root.winfo_screenheight()
    size = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
    print(size)
    root.geometry(size)
    canvas = tk.Canvas(root, width=300, height=300, highlightthickness=0, borderwidth=0)
    canvas.pack()
    img = []
    while 1:
        im = Image.open('load2.gif')
        # GIF图片流的迭代器
        iter = ImageSequence.Iterator(im)
        # frame就是gif的每一帧，转换一下格式就能显示了
        for frame in iter:
            pic = ImageTk.PhotoImage(frame)
            canvas.create_image((100, 150), image=pic)
            time.sleep(0.1)
            root.update_idletasks()  # 刷新
            root.update()


# 输入用户名密码窗体
class loginWindow(object):
    def __init__(self, window_msg):
        self.username = ""
        self.password = ""
        self.window_msg = window_msg

    def get_screen_ratio(self):
        hDC = win32gui.GetDC(0)
        dpiA = win32print.GetDeviceCaps(hDC, win32con.DESKTOPHORZRES) / win32print.GetDeviceCaps(hDC, win32con.HORZRES)
        dpiB = win32print.GetDeviceCaps(hDC, win32con.LOGPIXELSX) / 0.96 / 100
        if dpiA == 1:
            return dpiB
        elif dpiB == 1:
            return dpiA
        elif dpiA == dpiB:
            return dpiA

    def start_window(self):
        try:
            ratio = round(self.get_screen_ratio(), 2)
        except Exception:
            ratio = 1
        create_pic()
        print("ratio ", ratio)
        self.window = tk.Tk()
        print(self.window, 'self.window')
        self.window.iconbitmap("window_title.png")
        self.window.title('ChocLead')
        self.window.configure(bg='white')
        self.window.attributes('-topmost', True)
        self.var_usr_name = tk.StringVar()
        self.var_usr_pwd = tk.StringVar()
        self.var_lang = tk.Radiobutton()
        # window.geometry('400x200')
        width = 250 * ratio
        height = 150 * ratio
        screenwidth = self.window.winfo_screenwidth()
        screenheight = self.window.winfo_screenheight()
        alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
        print(alignstr)
        self.window.geometry(alignstr)
        self.window.resizable(0, 0)
        if self.window_msg != "":
            tk.Label(self.window, text=self.window_msg, fg="red", bg='white').place(x=85 * ratio, y=5 * ratio)
        tk.Label(self.window, text=_("username_des"), bg='white').place(x=10 * ratio, y=25 * ratio)
        tk.Label(self.window, text=_("Password: "), bg='white').place(x=10 * ratio, y=65 * ratio)
        # tk.Radiobutton(self.window, text=_("Password: "), bg='white').place(x=10 * ratio, y=105 * ratio)
        # 用户名输入框
        entry_usr_name = tk.Entry(self.window, textvariable=self.var_usr_name)
        entry_usr_name.place(x=85 * ratio, y=25 * ratio)
        # 密码输入框
        entry_usr_pwd = tk.Entry(self.window, textvariable=self.var_usr_pwd, show='*')
        entry_usr_pwd.place(x=85 * ratio, y=65 * ratio)
        # 登录 退出按钮
        bt_login = tk.Button(self.window, text=_("login_des"), command=self.usr_log_in)
        bt_login.place(x=65 * ratio, y=115 * ratio)
        bt_logquit = tk.Button(self.window, text=_("quit_des"), command=self.usr_sign_quit)
        bt_logquit.place(x=150 * ratio, y=115 * ratio)
        tk.Label(self.window, text=_("loginWarning"), fg="green", bg='white').place(x=30 * ratio, y=90 * ratio)
        # 主循环
        self.window.mainloop()
        delete_pic()
        user_dict = {"username": self.username, "password": self.password}
        # setLoginDict(user_dict)
        return user_dict

    # 登录函数
    def usr_log_in(self):
        # 从本地字典获取用户信息，如果没有则新建本地数据库
        # 判断用户名和密码是否匹配
        self.username = self.var_usr_name.get()
        self.password = self.var_usr_pwd.get()
        if self.username == '' or self.password == '':
            tk.messagebox.showerror(message=_("errorWarning"), bg='white')
        else:
            self.usr_sign_quit()

    # 退出的函数
    def usr_sign_quit(self):
        self.window.destroy()


# 默认签订失败，成功返回True
confirm_contract = False


def contract_window():
    """签署合同弹窗"""
    # 实例化object，创建窗口window
    # try:
    #     ratio = round(loginWindow.get_screen_ratio(), 2)
    # except Exception:
    #     ratio = 1
    window = tk.Tk()
    # 窗口名
    window.title('签署合同')
    window.geometry("200x200+-10+0")
    # 设定窗口大小（长*宽）
    width = 800
    height = 500
    screenwidth = window.winfo_screenwidth()
    screenheight = window.winfo_screenheight()
    alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
    window.geometry(alignstr)

    # window.resizable(0, 0)
    def text_contract():
        """解析合同内容"""
        document = Document('TecleadContract.docx')
        contract_data = ''
        for i in document.paragraphs:
            contract_data = f'{contract_data}\n{i.text}'
        return contract_data

    def to_confirm():
        """确定按钮点击事件：如果得到确定就返回True否则继续"""
        global confirm_contract
        if var1.get() == 1:
            confirm_contract = True
            window.destroy()
            return confirm_contract
        else:
            tkinter.messagebox.askokcancel(title='没有勾选哟', message='请勾选后点击确定')

    def to_exit():
        """退出确定"""
        exit = tkinter.messagebox.askokcancel(title='退出', message='您确定要退出吗？')
        if exit:
            window.destroy()

    var1 = tk.IntVar()  # 定义var1整型变量用来存放选择行为返回值
    l1 = tk.Label(window, text='请认真阅读以下文件', font=('Arial', 20, "bold"))  # Arial
    l1.pack()
    # 合同模块
    # 创建一个水平滚动条
    scrollbarl = Scrollbar(window, orient=HORIZONTAL)
    # 水平滚动条位于窗口底端，当窗口改变大小时会在X方向填满窗口
    scrollbarl.pack(side=BOTTOM, fill=X)
    # 创建一个垂直滚动条
    scrollbar2 = Scrollbar(window)
    # 垂直滚动条位于窗口右端，当窗口改变大小时会在Y方向填满窗口
    scrollbar2.pack(side=RIGHT, fill=Y)
    # 创建一个列表框， x方向的滚动条指 令是scrollbarl 对象的set()方法，
    # y方向的滚动条指令是scrollbar2对象的set()方法
    mylist = Listbox(window, xscrollcommand=scrollbarl.set, yscrollcommand=scrollbar2.set, height=20, width=240, bd=2)
    # 获取合同.docx中的数据
    text_data = text_contract()
    data_list = text_data.split('\n')
    for data in data_list:
        mylist.insert(END, data)
    # 列表框位于窗口左端，当窗口改变大小时会在X与Y方向填满窗口
    mylist.place(relx=0.5, rely=0.45, relheight=0.75, relwidth=0.9, anchor=CENTER)
    # 移动水平滚动条时,改变列表框的x方向可见范围
    scrollbarl.config(command=mylist.xview)
    # 移动垂直滚动条时，改变列表框的y方向可见范围
    scrollbar2.config(command=mylist.yview)
    # 签订协议
    c1 = tk.Checkbutton(window, text='我已认真阅读以上文件并且确认无误', variable=var1, onvalue=1, offvalue=0)  # 传值原理类似于radiobutton部件
    c1.place(relx=0.5, rely=0.85, relheight=0.03, relwidth=0.29, anchor=CENTER)
    # 确定和取消按钮
    b1 = tk.Button(window, text='确定', relief=GROOVE, font=('Arial', 10), command=to_confirm, width=8,
                   height=1)  # width=8, height=1
    b2 = tk.Button(window, text='取消', relief=GROOVE, font=('Arial', 10), command=to_exit, width=8, height=1)  #
    b1.place(relx=0.4, rely=0.92, relheight=0.06, relwidth=0.1, anchor=CENTER)
    b2.place(relx=0.6, rely=0.92, relheight=0.06, relwidth=0.1, anchor=CENTER)
    # 主窗口循环
    window.mainloop()
    if confirm_contract:
        return True


def log_window():
    """日志分析器"""

    def to_exit():
        """退出确定"""
        exit = tkinter.messagebox.askokcancel(title='退出', message='您确定要退出吗？')
        if exit:
            window.destroy()

    def text_contract():
        """解析合同内容"""
        document = Document('TecleadContract.docx')
        contract_data = ''
        for i in document.paragraphs:
            contract_data = f'{contract_data}\n{i.text}'
        return contract_data

    def go(*args):  # 处理事件，*args表示可变参数
        print(comboxlist.get())  # 打印选中的值

    window = tk.Tk()

    def show():
        # global mylist
        # mylist.SetCaretIndex(mylist.GetCount() - 1, TRUE)
        print("作品:<< % s>>" % e1.get())
        print("作者:<< % s>>" % e2.get())
        e1.delete(0, END)
        e2.delete(0, END)

    # 窗口名
    window.title('Log Parser')
    window.geometry("200x200+-10+0")
    # 设定窗口大小（长*宽）
    width = 800
    height = 500
    screenwidth = window.winfo_screenwidth()
    screenheight = window.winfo_screenheight()
    alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
    window.geometry(alignstr)

    # title
    menubar = tk.Menu(window)
    # 创建一个File菜单项（默认不下拉，下拉内容包括New，Open，Save，Exit功能项）
    log_menu = tk.Menu(menubar, tearoff=0)

    # 将上面定义的空菜单命名为log_menu，放在菜单栏中，就是装入那个容器中
    menubar.add_cascade(label='Log', menu=log_menu)

    # 在Log中加入New、Open、Save等小菜单，即我们平时看到的下拉菜单，每一个小菜单对应命令操作.

    log_menu.add_command(label='New', command=to_exit)
    log_menu.add_command(label='Open', command=to_exit)
    log_menu.add_command(label='Save', command=to_exit)
    log_menu.add_separator()  # 添加一条分隔线
    log_menu.add_command(label='Add', command=window.quit)
    log_menu.add_command(label='Delete', command=window.quit)
    log_menu.add_separator()  # 添加一条分隔线
    log_menu.add_command(label='Exit', command=window.quit)

    # 第7步，创建一个Edit菜单项（默认不下拉，下拉内容包括Cut，Copy，Paste功能项）
    editmenu = tk.Menu(menubar, tearoff=0)
    # 将上面定义的空菜单命名为 Edit，放在菜单栏中，就是装入那个容器中
    menubar.add_cascade(label='Edit', menu=editmenu)

    # 同样的在 Edit 中加入Cut、Copy、Paste等小命令功能单元，如果点击这些单元, 就会触发to_exit的功能
    editmenu.add_command(label='Cut', command=to_exit)
    editmenu.add_command(label='Copy', command=to_exit)
    editmenu.add_command(label='Paste', command=to_exit)

    # 让主菜单显示出来
    window.config(menu=menubar)

    Label(window, text="").grid(row=0, column=0, padx=10)
    Label(window, text="起始时间:").grid(row=0, column=1)
    Label(window, text="截至时间:").grid(row=1, column=1)
    Label(window, text="").grid(row=0, column=3, padx=50)
    Label(window, text="指定用户:").grid(row=0, column=4)
    Label(window, text="搜索类型:").grid(row=1, column=4)
    e1 = Entry(window)
    e2 = Entry(window)
    e3 = Entry(window)
    # 第一行第一列
    e1.grid(row=0, column=2, padx=10, pady=5)
    # 第二行第一列
    e2.grid(row=1, column=2, padx=10, pady=5)
    e3.grid(row=0, column=5, padx=10, pady=5)
    # 选择框
    comvalue = tkinter.StringVar()  # 窗体自带的文本，新建一个值
    comboxlist = ttk.Combobox(window, textvariable=comvalue, width=18)  # 初始化
    comboxlist["values"] = ("创建流程", "更改流程", "删除流程", "运行流程", "其他")

    comboxlist.current(0)  # 选择第一个
    comboxlist.bind("<<ComboboxSelected>>", go)
    comboxlist.grid(row=1, column=5, padx=10, pady=5)
    # 如果表格大于组件，那么可以使用sticky选项来设置组件的位置
    # 同样你需要使用N，E，S,W以及他们的组合NE，SE，SW，NW来表示方位
    Label(window, text="").grid(row=0, column=6, padx=20)
    Button(window, text="获取信息", width=10, command=show).grid(row=0, column=7, sticky=W, padx=10, pady=5)
    Button(window, text="退出", width=10, command=to_exit).grid(row=1, column=7, sticky=E, padx=10, pady=5)

    # 创建一个水平滚动条
    scrollbarl = Scrollbar(window, orient=HORIZONTAL)
    # 水平滚动条位于窗口底端，当窗口改变大小时会在X方向填满窗口
    # Label(window, text="").grid(row=3, column=0, padx=10)
    # scrollbarl.grid(row=3, column=1, sticky=S, pady=310)
    scrollbarl.place(relx=0.025, rely=0.97, relheight=0.025, relwidth=0.95)
    # 创建一个垂直滚动条
    scrollbar2 = Scrollbar(window)
    # 垂直滚动条位于窗口右端，当窗口改变大小时会在Y方向填满窗口
    # scrollbar2.grid(row=2, column=0, sticky=E)
    scrollbar2.place(relx=0.98, rely=0.1, relheight=0.9, relwidth=0.018)
    # 创建一个列表框， x方向的滚动条指 令是scrollbarl 对象的set()方法，
    # y方向的滚动条指令是scrollbar2对象的set()方法
    mylist = Listbox(window, xscrollcommand=scrollbarl.set, yscrollcommand=scrollbar2.set, height=10, width=100,
                     bd=2)  # , height=20, width=240, bd=2
    # 获取合同.docx中的数据
    text_data = text_contract()
    data_list = text_data.split('\n')
    for data in data_list:
        mylist.insert(END, data)
    # 列表框位于窗口左端，当窗口改变大小时会在X与Y方向填满窗口
    mylist.place(relx=0.5, rely=0.55, relheight=0.75, relwidth=0.92, anchor=CENTER)
    # mylist.grid(row=2, column=1, rowspan=1, columnspan=100)
    # 移动水平滚动条时,改变列表框的x方向可见范围
    scrollbarl.config(command=mylist.xview)
    # 移动垂直滚动条时，改变列表框的y方向可见范围
    scrollbar2.config(command=mylist.yview)

    # 主窗口循环
    window.mainloop()
# if __name__ == '__main__':
# contract_window()
# if confirm_contract:
#     print('收到了')
# else:
# #     print('无')
# log_window()
